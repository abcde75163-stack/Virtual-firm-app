"""
docx_builder.py  v2.0
마크다운 형식의 VF TBD 보고서 텍스트를 python-docx로 Word 문서 변환

개선사항:
- 여백 "좁게" (상하좌우 1.27cm)
- 표지 + 목차 페이지
- 장(#) 시작 시 페이지 강제 이동
- 표 제목 [표 ...] 패턴 처리
- **볼드** 인라인 처리
- ❍ / □ / ※ 패턴별 스타일 처리
- 빈 줄로 문단 간격 표현
- 마크다운 표 → 헤더 네이비 + 행 교대 배경 스타일 표
"""
import io
import re
import zipfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 색상 ──────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x38, 0x64)
BLUE   = RGBColor(0x2E, 0x75, 0xB6)
LLBLUE = RGBColor(0xDE, 0xEA, 0xF1)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xF2, 0xF2, 0xF2)
ALTROW = RGBColor(0xF8, 0xF9, 0xFA)
BLACK  = RGBColor(0x00, 0x00, 0x00)
DGRAY  = RGBColor(0x40, 0x40, 0x40)
MGRAY  = RGBColor(0x66, 0x66, 0x66)
RED    = RGBColor(0xC0, 0x00, 0x00)

# ── 여백 "좁게" = 1.27cm ─────────────────────────────────
MARGIN_NARROW = Cm(1.27)


# ════════════════════════════════════════════════════════
# XML 헬퍼
# ════════════════════════════════════════════════════════
def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def _set_cell_border(cell, color="BBBBBB", size="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    # tcBorders는 shd(배경) 이전에 삽입해야 스키마 순서를 지킬 수 있음
    shd_elem = tcPr.find(qn("w:shd"))
    if shd_elem is not None:
        tcPr.insert(list(tcPr).index(shd_elem), tcBorders)
    else:
        tcPr.append(tcBorders)

def _set_cell_valign(cell, align="center"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)

def _add_border_bottom(para, color="2E75B6", size=12):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def _add_page_break_before(para):
    pPr = para._p.get_or_add_pPr()
    pb  = OxmlElement("w:pageBreakBefore")
    pb.set(qn("w:val"), "true")
    pPr.append(pb)


# ════════════════════════════════════════════════════════
# 인라인 볼드 처리 헬퍼
# ════════════════════════════════════════════════════════
def _add_runs_with_bold(para, text: str, base_size: Pt, base_color: RGBColor,
                         font_name="맑은 고딕", base_bold=False):
    """**볼드** 패턴을 인식해서 run을 분리 추가"""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            run = para.add_run(inner)
            run.font.name  = font_name
            run.font.size  = base_size
            run.font.bold  = True
            run.font.color.rgb = base_color
        else:
            run = para.add_run(part)
            run.font.name  = font_name
            run.font.size  = base_size
            run.font.bold  = base_bold
            run.font.color.rgb = base_color


# ════════════════════════════════════════════════════════
# 표지 생성
# ════════════════════════════════════════════════════════
def _add_cover(doc: Document, patent_name: str, demand_company: str):
    for _ in range(5):
        doc.add_paragraph()

    # Virtual Firm
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Virtual Firm")
    r.font.name = "맑은 고딕"; r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tech Business Development 보고서")
    r.font.name = "맑은 고딕"; r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = BLUE

    # 구분선
    div = doc.add_paragraph()
    _add_border_bottom(div, "1F3864", 12)
    div.paragraph_format.space_before = Pt(8)
    div.paragraph_format.space_after  = Pt(12)

    # 기술명
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(patent_name)
    r.font.name = "맑은 고딕"; r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = NAVY

    if demand_company:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"수요기업: {demand_company}")
        r.font.name = "맑은 고딕"; r.font.size = Pt(12); r.font.color.rgb = DGRAY

    for _ in range(7):
        doc.add_paragraph()

    now = datetime.now().strftime("%Y년 %m월")
    for text, size, color in [
        ("부산대학교 산학협력단 기술사업화팀 (TLO)", 11, DGRAY),
        (f"보고서 버전: VF TBD v3.0  |  작성일: {now}", 10, MGRAY),
        ("※ 본 보고서는 참고용 자료입니다.", 9, MGRAY),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "맑은 고딕"; r.font.size = Pt(size); r.font.color.rgb = color


# ════════════════════════════════════════════════════════
# 목차 생성
# ════════════════════════════════════════════════════════
TOC_ENTRIES = [
    ("Ⅰ. 기술 개요", True),
    ("  Ⅰ-1. 기술명 및 기술 정의", False),
    ("  Ⅰ-2. 기술 배경 및 산업적 필요성", False),
    ("  Ⅰ-3. 특허 명세서 기반 핵심 구성 및 작동 원리", False),
    ("  Ⅰ-4. 주요 실시예 및 구현 방식", False),
    ("  Ⅰ-5. 청구항 기반 권리범위 분석", False),
    ("  Ⅰ-6. 기술적 효과 및 경쟁 우위", False),
    ("Ⅱ. 문제점 및 해결방안", True),
    ("  Ⅱ-1. 기존 시장·기술의 문제점 및 Pain Point", False),
    ("  Ⅱ-2. 기존 대안 및 한계", False),
    ("  Ⅱ-3. 해결방안 및 기술-고객 가치 연결", False),
    ("  Ⅱ-4. 경쟁 포지셔닝 및 도입 필요성", False),
    ("  Ⅱ-5. 사업화 리스크 및 대응방안", False),
    ("  Ⅱ-6. PSST 기반 Problem-Solution 정리", False),
    ("Ⅲ. Scale-up 전략", True),
    ("  Ⅲ-1. 시장 진입 가능성 및 세분화 (TAM-SAM-SOM)", False),
    ("  Ⅲ-2. 타깃 고객 및 수요기업 유형", False),
    ("  Ⅲ-3. 사업화 모델 및 경로", False),
    ("  Ⅲ-4. 제품화·서비스화 전략", False),
    ("  Ⅲ-5. IP 고도화 및 인증·규제 검토", False),
    ("  Ⅲ-6. 실증·PoC 및 단계별 로드맵", False),
    ("  Ⅲ-7. 수익모델·기술료 및 재무추정", False),
    ("  Ⅲ-8. Scale-up 핵심 성공요인", False),
    ("Ⅳ. Virtual Firm 활용", True),
    ("  Ⅳ-1. Virtual Firm 개요", False),
    ("  Ⅳ-2. Lean Canvas 및 BMC", False),
    ("  Ⅳ-3. SWOT 분석", False),
    ("  Ⅳ-4. STP 및 GTM 전략", False),
    ("  Ⅳ-5. 연구자·기술보유자 역량", False),
    ("  Ⅳ-6. 사업타당성 종합 판단 및 결론", False),
]

def _add_toc(doc: Document):
    # 목차 제목
    p = doc.add_paragraph()
    _add_border_bottom(p, "1F3864", 12)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run("목  차")
    r.font.name = "맑은 고딕"; r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = NAVY

    for entry, is_chapter in TOC_ENTRIES:
        p = doc.add_paragraph()
        if is_chapter:
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(entry.strip())
            r.font.name = "맑은 고딕"; r.font.size = Pt(11)
            r.font.bold = True; r.font.color.rgb = NAVY
        else:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Cm(0.8)
            r = p.add_run(entry.strip())
            r.font.name = "맑은 고딕"; r.font.size = Pt(10)
            r.font.bold = False; r.font.color.rgb = DGRAY

    doc.add_page_break()


# ════════════════════════════════════════════════════════
# 마크다운 표 → Word 표
# ════════════════════════════════════════════════════════
def _parse_md_table(lines: list) -> list:
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # 구분선 행 (|---|---|) 스킵
        inner = stripped.strip("|").replace(" ", "")
        if all(c in "-:|" for c in inner) and "-" in inner:
            continue
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        rows.append(cells)
    return rows

def _add_md_table(doc: Document, table_lines: list):
    rows = _parse_md_table(table_lines)
    if len(rows) < 2:
        return

    col_count = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    # 열 너비 균등 분배 (A4 좁게 기준 content width ≈ 10466 DXA)
    CONT_DXA = 10466
    col_w = CONT_DXA // col_count
    for row_idx, row_data in enumerate(rows):
        is_header = (row_idx == 0)
        row = tbl.rows[row_idx]
        # 행 높이 최소값
        row.height = Pt(18)

        for col_idx in range(col_count):
            cell = row.cells[col_idx]
            text = row_data[col_idx] if col_idx < len(row_data) else ""

            # 배경색
            if is_header:
                _set_cell_bg(cell, "1F3864")
                _set_cell_border(cell, "2E75B6", "6")
            elif row_idx % 2 == 1:
                _set_cell_bg(cell, "F8F9FA")
                _set_cell_border(cell, "CCCCCC", "4")
            else:
                _set_cell_bg(cell, "FFFFFF")
                _set_cell_border(cell, "CCCCCC", "4")

            _set_cell_valign(cell, "center")

            # 텍스트
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.15)

            txt = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  # 볼드 태그 제거 (셀은 단순 처리)
            run = p.add_run(txt)
            run.font.name  = "맑은 고딕"
            run.font.size  = Pt(9)
            run.font.bold  = is_header
            run.font.color.rgb = WHITE if is_header else BLACK

    # 표 뒤 여백
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════════════
# 메인 변환 함수
# ════════════════════════════════════════════════════════
def build_docx(
    report_text:    str,
    patent_name:    str = "특허기술",
    demand_company: str = "",
) -> bytes:
    doc = Document()

    # ── 여백 "좁게" 적용 ─────────────────────────────────
    for sec in doc.sections:
        sec.page_width    = Cm(21.0)
        sec.page_height   = Cm(29.7)
        sec.top_margin    = MARGIN_NARROW
        sec.bottom_margin = MARGIN_NARROW
        sec.left_margin   = MARGIN_NARROW
        sec.right_margin  = MARGIN_NARROW

    # ── 기본 폰트 ─────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    style.font.color.rgb = BLACK

    # ── 표지 ─────────────────────────────────────────────
    _add_cover(doc, patent_name, demand_company)
    doc.add_page_break()

    # ── 목차 ─────────────────────────────────────────────
    _add_toc(doc)

    # ── 본문 파싱 ─────────────────────────────────────────
    lines = report_text.split("\n")
    i, n  = 0, len(lines)

    # 장 번호 패턴 (# 으로 시작하는 장 제목)
    CHAPTER_RE = re.compile(r"^#\s+(Ⅰ|Ⅱ|Ⅲ|Ⅳ)\.")
    # 섹션 제목 패턴 (## Ⅰ-1 등)
    SECTION_RE = re.compile(r"^##\s+")
    # 서브섹션 패턴 (### )
    SUBSEC_RE  = re.compile(r"^###\s+")
    # 표 제목 패턴 [표 ...]
    TBLTTL_RE  = re.compile(r"^\[표\s")
    # 불릿 패턴
    BULLET_RE  = re.compile(r"^\s*[-•]\s+")
    # 번호 불릿 (1. 2. 등)
    NUMLIST_RE = re.compile(r"^\s*\d+\.\s+")
    # 구분선
    DIVIDER_RE = re.compile(r"^-{3,}$")
    # ※ 주석
    NOTE_RE    = re.compile(r"^[※\*]")
    # ❍ 서브헤더
    SUB_RE     = re.compile(r"^[❍◎□]\s*")

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # ── 빈 줄: 완전 무시 (space_before/after로 간격 제어)
        if not stripped:
            i += 1
            continue

        # ── 마크다운 표 감지 ─────────────────────────────────
        if stripped.startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_md_table(doc, table_lines)
            continue

        # ── 장 제목 # (새 페이지 강제) ──────────────────────
        if CHAPTER_RE.match(stripped):
            text = re.sub(r"^#+\s+", "", stripped)
            p = doc.add_paragraph()
            _add_page_break_before(p)
            _add_border_bottom(p, "1F3864", 16)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(10)
            _add_runs_with_bold(p, text, Pt(15), NAVY, base_bold=True)
            i += 1
            continue

        # ── 섹션 제목 ## ────────────────────────────────────
        if SECTION_RE.match(stripped):
            text = re.sub(r"^#+\s+", "", stripped)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(5)
            r0 = p.add_run("■ ")
            r0.font.name = "맑은 고딕"; r0.font.size = Pt(11.5)
            r0.font.bold = True; r0.font.color.rgb = BLUE
            _add_runs_with_bold(p, text, Pt(11.5), NAVY, base_bold=True)
            i += 1
            continue

        # ── 서브섹션 ### ────────────────────────────────────
        if SUBSEC_RE.match(stripped):
            text = re.sub(r"^#+\s+", "", stripped)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)
            r0 = p.add_run("◎ ")
            r0.font.name = "맑은 고딕"; r0.font.size = Pt(10.5)
            r0.font.bold = True; r0.font.color.rgb = BLUE
            _add_runs_with_bold(p, text, Pt(10.5), DGRAY, base_bold=True)
            i += 1
            continue

        # ── 표 제목 [표 ...] ────────────────────────────────
        if TBLTTL_RE.match(stripped):
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(text)
            run.font.name = "맑은 고딕"; run.font.size = Pt(10)
            run.font.bold = True; run.font.color.rgb = NAVY
            i += 1
            continue

        # ── ❍ / □ 서브헤더 ───────────────────────────────
        if SUB_RE.match(stripped):
            text = SUB_RE.sub("", stripped)
            prefix = stripped[0]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(3)
            r0 = p.add_run(prefix + " ")
            r0.font.name = "맑은 고딕"; r0.font.size = Pt(10.5)
            r0.font.bold = True; r0.font.color.rgb = BLUE
            _add_runs_with_bold(p, text, Pt(10.5), DGRAY, base_bold=True)
            i += 1
            continue

        # ── 불릿 - ───────────────────────────────────────
        if BULLET_RE.match(line):
            text = BULLET_RE.sub("", stripped)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            r0 = p.add_run("• ")
            r0.font.name = "맑은 고딕"; r0.font.size = Pt(9.5); r0.font.color.rgb = BLUE
            _add_runs_with_bold(p, text, Pt(9.5), BLACK)
            i += 1
            continue

        # ── 번호 불릿 1. 2. ─────────────────────────────
        if NUMLIST_RE.match(stripped):
            text = NUMLIST_RE.sub("", stripped)
            match = NUMLIST_RE.match(stripped)
            num_str = match.group(0).strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            r0 = p.add_run(num_str + " ")
            r0.font.name = "맑은 고딕"; r0.font.size = Pt(9.5)
            r0.font.bold = True; r0.font.color.rgb = NAVY
            _add_runs_with_bold(p, text, Pt(9.5), BLACK)
            i += 1
            continue

        # ── 구분선 --- ─────────────────────────────────
        if DIVIDER_RE.match(stripped):
            p = doc.add_paragraph()
            _add_border_bottom(p, "CCCCCC", 4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            i += 1
            continue

        # ── ※ 주석 ───────────────────────────────────────
        if NOTE_RE.match(stripped):
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(5)
            p.paragraph_format.left_indent  = Cm(0.3)
            run = p.add_run(text)
            run.font.name = "맑은 고딕"; run.font.size = Pt(9)
            run.font.color.rgb = MGRAY
            i += 1
            continue

        # ── 일반 텍스트 ──────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        _add_runs_with_bold(p, stripped, Pt(10), BLACK)
        i += 1

    # ── 푸터 ──────────────────────────────────────────────
    _add_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return _fix_docx(buf.getvalue())


def _fix_docx(docx_bytes: bytes) -> bytes:
    """python-docx 생성 시 발생하는 settings.xml zoom 속성 누락 오류 수정"""
    buf_in  = io.BytesIO(docx_bytes)
    buf_out = io.BytesIO()
    with zipfile.ZipFile(buf_in, 'r') as zin:
        with zipfile.ZipFile(buf_out, 'w', zipfile.ZIP_DEFLATED) as zout:
            for fname in zin.namelist():
                data = zin.read(fname)
                if fname == 'word/settings.xml':
                    content = data.decode('utf-8')
                    content = re.sub(
                        r'<w:zoom([^/]*?)/>',
                        lambda m: f'<w:zoom{m.group(1)} w:percent="100"/>',
                        content
                    )
                    data = content.encode('utf-8')
                zout.writestr(fname, data)
    return buf_out.getvalue()


def _add_footer(doc: Document):
    for sec in doc.sections:
        footer_p = sec.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run(
            "부산대학교 산학협력단 TLO  |  VF TBD Report v3.0  |  본 보고서는 참고용 자료입니다"
        )
        run.font.name = "맑은 고딕"
        run.font.size = Pt(8)
        run.font.color.rgb = MGRAY
