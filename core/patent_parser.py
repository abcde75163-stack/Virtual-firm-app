"""
patent_parser.py
PDF / DOCX 파일에서 텍스트를 추출하는 모듈
"""
import os
from pathlib import Path


def extract_text_from_file(file_path: str) -> str:
    """
    파일 경로를 받아 텍스트를 추출하여 반환
    지원 형식: PDF, DOCX, TXT
    """
    suffix = Path(file_path).suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(file_path)
    elif suffix == ".docx":
        return _extract_docx(file_path)
    elif suffix == ".txt":
        return _extract_txt(file_path)
    else:
        raise ValueError(f"지원하지 않는 파일 형식입니다: {suffix}")


def _extract_pdf(path: str) -> str:
    """PyMuPDF(fitz)로 PDF 텍스트 추출"""
    try:
        import fitz  # PyMuPDF
        doc  = fitz.open(path)
        text = ""
        for page_num, page in enumerate(doc):
            page_text = page.get_text("text")
            if page_text.strip():
                text += f"\n[페이지 {page_num + 1}]\n{page_text}"
        doc.close()

        text = text.strip()
        if not text:
            raise ValueError("PDF에서 텍스트를 추출할 수 없습니다. 스캔된 이미지 PDF일 수 있습니다.")
        return text

    except ImportError:
        raise ImportError(
            "PyMuPDF 라이브러리가 없습니다. "
            "pip install pymupdf 로 설치해 주세요."
        )


def _extract_docx(path: str) -> str:
    """python-docx로 DOCX 텍스트 추출"""
    try:
        from docx import Document
        import re

        doc   = Document(path)
        lines = []

        # 본문 단락
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                lines.append(t)

        # 표 내용
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    lines.append(" | ".join(row_texts))

        text = "\n".join(lines)
        if not text.strip():
            raise ValueError("DOCX에서 텍스트를 추출할 수 없습니다. 파일을 확인해 주세요.")
        return text

    except ImportError:
        raise ImportError(
            "python-docx 라이브러리가 없습니다. "
            "pip install python-docx 로 설치해 주세요."
        )


def _extract_txt(path: str) -> str:
    """TXT 파일 읽기 (UTF-8 / EUC-KR 자동 감지)"""
    for encoding in ("utf-8", "euc-kr", "cp949", "utf-8-sig"):
        try:
            with open(path, "r", encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError("TXT 파일의 인코딩을 인식할 수 없습니다.")


def truncate_text(text: str, max_chars: int = 30000) -> str:
    """
    토큰 한도 초과 방지를 위한 텍스트 트리밍
    명세서 앞부분(발명의 명칭~청구항)을 최대한 보존
    """
    if len(text) <= max_chars:
        return text

    # 청구항 위치 탐색 — 청구항 이후를 포함하도록 우선 보존
    claim_idx = -1
    for keyword in ["【특허청구범위】", "【청구항", "특허청구범위", "청구항 1"]:
        idx = text.find(keyword)
        if idx != -1:
            claim_idx = idx
            break

    if claim_idx != -1 and claim_idx < max_chars:
        # 청구항이 max_chars 내에 있으면 그냥 자름
        return text[:max_chars] + "\n\n[※ 명세서가 길어 일부를 생략하였습니다]"
    else:
        # 앞 2/3 + 청구항 영역
        front = text[: int(max_chars * 0.7)]
        back  = text[max(0, len(text) - int(max_chars * 0.3)) :]
        return front + "\n\n[...중략...]\n\n" + back + "\n\n[※ 명세서가 길어 일부를 생략하였습니다]"
